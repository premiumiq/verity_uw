"""One-shot script to produce a starter .docx template for a report.

Usage (from the project root, via venv):
    python -m verity.reporting._template_authoring model_inventory

This generates a .docx file at:
    verity/src/verity/reporting/templates/<report_code>.docx

The output is a *starter* — a styled Word document with the correct
docxtpl placeholders ({{...}}, {%tr ... %}). Designers/compliance teams
open this in Word and restyle freely (logo, colors, watermark) without
touching placeholders.

Why a script: the .docx files in the repo are binary blobs. Generating
them from code keeps the placeholder list documented and makes
starter-content changes reviewable in PRs.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document  # type: ignore
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches, Cm


TEMPLATES_DIR = Path(__file__).parent / "templates"


# =============================================================================
# Styling constants — change these in one place and rerun the script.
# =============================================================================

HEADING_FONT       = "Georgia"
BODY_FONT          = "Poppins"
BRAND_BLUE         = RGBColor(0x1A, 0x36, 0x5D)
BRAND_BLUE_LIGHT   = RGBColor(0x6E, 0x8F, 0xCC)
BODY_GRAY          = RGBColor(0x33, 0x33, 0x33)
MUTED_GRAY         = RGBColor(0x6B, 0x6B, 0x6B)
CARD_SHADE_HEX     = "F4F6FA"   # subtle cool gray for card cell background


# =============================================================================
# Low-level helpers
# =============================================================================

def _set_run_font(run, *, name=BODY_FONT, size=11, bold=False, italic=False, color=BODY_GRAY):
    run.font.name = name
    # `w:eastAsia` font also needs to be set for some renderers.
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"),    name)
    rFonts.set(qn("w:hAnsi"),    name)
    rFonts.set(qn("w:cs"),       name)
    rFonts.set(qn("w:eastAsia"), name)
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    run.font.color.rgb = color


def _add_styled_para(doc, text="", *, font=BODY_FONT, size=11, bold=False,
                     italic=False, color=BODY_GRAY, align=WD_ALIGN_PARAGRAPH.LEFT,
                     space_before=0, space_after=4):
    """Add a paragraph with a single styled run."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        _set_run_font(run, name=font, size=size, bold=bold, italic=italic, color=color)
    return p


def _add_heading(doc, text, *, level=1, page_break_before=False):
    """Add a heading paragraph using Georgia + brand blue.

    level=1 → 22pt, level=2 → 16pt, level=3 → 13pt.
    """
    sizes = {1: 22, 2: 16, 3: 13}
    p = doc.add_paragraph()
    if page_break_before:
        p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(6 if level > 1 else 0)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    _set_run_font(run, name=HEADING_FONT, size=sizes[level], bold=True, color=BRAND_BLUE)
    return p


def _add_page_break(doc):
    """Insert a hard page break."""
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _shade_cell(cell, hex_color):
    """Apply background shading to a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell, *, color="C5CDDC", size_eighths=6):
    """Apply 1pt borders to all four sides of a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    str(size_eighths))
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _enable_different_first_page_header_footer(section):
    """Tell Word: page 1 of this section gets its own (empty) header/footer."""
    sectPr = section._sectPr
    titlePg = sectPr.find(qn("w:titlePg"))
    if titlePg is None:
        titlePg = OxmlElement("w:titlePg")
        sectPr.append(titlePg)


# =============================================================================
# Card builders — each asset type's section uses a single-column table where
# every body row is a "card" repeated by docxtpl via {%tr%}.
# =============================================================================

def _add_asset_card_table(doc, *, item_var: str, list_var: str):
    """Add a single-column table that docxtpl will repeat once per asset.

    Each card cell contains:
      - Title line: {{ display_name }}  v{{ version_label }}  ·  {{ status }}
      - Description (italic, smaller, gray)
      - Facts line: Materiality · Owner · Application(s)
    """
    # Three rows: {%tr open%} · body · {%tr close%}.
    table = doc.add_table(rows=3, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)

    # tr-open marker row
    tr_open = table.rows[0].cells[0]
    tr_open.text = ""
    p = tr_open.paragraphs[0]
    p.add_run("{%tr for a in " + list_var + " %}")

    # body row — the actual card
    body = table.rows[1].cells[0]
    body.width = Inches(6.5)
    _shade_cell(body, CARD_SHADE_HEX)
    _set_cell_borders(body, color="C5CDDC")
    body.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    # Clear the default empty paragraph
    body._tc.remove(body.paragraphs[0]._p)

    # Card paragraph 1 — title line
    p = body.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run("{{ a.display_name }}")
    _set_run_font(run, name=BODY_FONT, size=13, bold=True, color=BRAND_BLUE)
    run = p.add_run("    v{{ a.version_label }}    ·    ")
    _set_run_font(run, name=BODY_FONT, size=11, color=MUTED_GRAY)
    run = p.add_run("{{ a.lifecycle_state_display }}")
    _set_run_font(run, name=BODY_FONT, size=11, bold=True, color=BRAND_BLUE_LIGHT)

    # Card paragraph 2 — description (italic, gray)
    p = body.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("{{ a.entity_description or '—' }}")
    _set_run_font(run, name=BODY_FONT, size=10, italic=True, color=MUTED_GRAY)

    # Card paragraph 3 — facts (Materiality · Owner · Apps)
    p = body.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)
    facts = [
        ("Materiality: ",  "{{ a.materiality_display }}"),
        ("    ·    Owner: ",         "{{ a.owner_name or '—' }}"),
        ("    ·    Application(s): ", "{{ a.applications }}"),
    ]
    for label, value in facts:
        run = p.add_run(label)
        _set_run_font(run, name=BODY_FONT, size=10, color=MUTED_GRAY)
        run = p.add_run(value)
        _set_run_font(run, name=BODY_FONT, size=10, bold=True, color=BODY_GRAY)

    # tr-close marker row
    tr_close = table.rows[2].cells[0]
    tr_close.text = ""
    p = tr_close.paragraphs[0]
    p.add_run("{%tr endfor %}")
    return table


def _add_lifecycle_table(doc, list_var: str):
    """Add a 5-column lifecycle-events table for one asset type."""
    table = doc.add_table(rows=4, cols=5)
    table.style = "Light Grid Accent 1"

    headers = ["Asset", "Transition", "Gate", "Approver", "Approved At"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        _set_run_font(run, name=BODY_FONT, size=10, bold=True, color=BRAND_BLUE)

    # tr-open
    tr_open = table.rows[1].cells[0]
    tr_open.text = "{%tr for ev in " + list_var + " %}"
    table.rows[1].cells[0].merge(table.rows[1].cells[4])

    # body
    body = table.rows[2].cells
    cells_text = [
        "{{ ev.asset_display_name }}  v{{ ev.version_label }}",
        "{{ ev.from_state_display }} → {{ ev.to_state_display }}",
        "{{ ev.gate_type }}",
        "{{ ev.approver_name }}",
        "{{ ev.approved_at.strftime('%Y-%m-%d %H:%M') if ev.approved_at else '—' }}",
    ]
    for i, txt in enumerate(cells_text):
        body[i].text = ""
        p = body[i].paragraphs[0]
        run = p.add_run(txt)
        _set_run_font(run, name=BODY_FONT, size=10, color=BODY_GRAY)

    # tr-close
    tr_close = table.rows[3].cells[0]
    tr_close.text = "{%tr endfor %}"
    table.rows[3].cells[0].merge(table.rows[3].cells[4])
    return table


# =============================================================================
# Model Inventory starter template
# =============================================================================

def author_model_inventory() -> Path:
    doc = Document()

    # Establish Normal style font (Poppins) globally.
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    rPr = normal.element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), BODY_FONT)
    rFonts.set(qn("w:hAnsi"), BODY_FONT)
    rFonts.set(qn("w:cs"),    BODY_FONT)
    rFonts.set(qn("w:eastAsia"), BODY_FONT)
    # remove any existing rFonts before appending
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rPr.append(rFonts)

    # ── Page setup ─────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin    = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin   = Inches(0.9)
    section.right_margin  = Inches(0.9)
    _enable_different_first_page_header_footer(section)

    # The first-page header/footer are explicitly empty (no chrome on cover).
    fp_header = section.first_page_header
    fp_header.is_linked_to_previous = False
    fp_header.paragraphs[0].text = ""
    fp_footer = section.first_page_footer
    fp_footer.is_linked_to_previous = False
    fp_footer.paragraphs[0].text = ""

    # Default header (page 2+)
    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_p.add_run("VERITY · {{ report_title }}")
    _set_run_font(run, name=HEADING_FONT, size=9, bold=True, color=BRAND_BLUE)

    # Default footer (page 2+)
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = footer_p.add_run("{{ report_title }}  ·  Generated {{ generated_at_str }}  ·  Scope: {{ scope_summary }}")
    _set_run_font(run, name=BODY_FONT, size=8, color=MUTED_GRAY)

    # ════════════════════════════════════════════════════════════
    # PAGE 1 — COVER
    # ════════════════════════════════════════════════════════════
    # Vertical spacing to push title to roughly center of page.
    for _ in range(8):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("VERITY")
    _set_run_font(run, name=HEADING_FONT, size=14, bold=True, color=BRAND_BLUE_LIGHT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Compliance Report")
    _set_run_font(run, name=HEADING_FONT, size=11, color=MUTED_GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("{{ report_title }}")
    _set_run_font(run, name=HEADING_FONT, size=36, bold=True, color=BRAND_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run("Generated {{ generated_at_str }}")
    _set_run_font(run, name=BODY_FONT, size=12, color=BODY_GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run("Scope: {{ scope_summary }}")
    _set_run_font(run, name=BODY_FONT, size=11, italic=True, color=MUTED_GRAY)

    _add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # PAGE 2 — PURPOSE & AUDIENCE
    # ════════════════════════════════════════════════════════════
    _add_heading(doc, "Purpose and Audience", level=1)

    _add_styled_para(
        doc,
        "Purpose",
        font=HEADING_FONT, size=13, bold=True, color=BRAND_BLUE,
        space_before=10, space_after=4,
    )
    _add_styled_para(
        doc,
        "{{ report_description }}",
        font=BODY_FONT, size=11, color=BODY_GRAY, space_after=10,
    )
    _add_styled_para(
        doc,
        "This report enumerates {{ total_count }} AI asset version(s) registered to "
        "Verity at the time of generation: {{ agent_count }} agent(s), "
        "{{ task_count }} task(s), and {{ prompt_count }} prompt(s). "
        "The breakdown by materiality is {{ high_count }} high, {{ medium_count }} medium, "
        "and {{ low_count }} low.",
        font=BODY_FONT, size=11, color=BODY_GRAY, space_after=14,
    )

    _add_styled_para(
        doc,
        "Intended audience",
        font=HEADING_FONT, size=13, bold=True, color=BRAND_BLUE,
        space_before=10, space_after=4,
    )
    _add_styled_para(
        doc,
        "Model Risk Management committees, compliance officers, internal audit teams, "
        "and external regulatory examiners. The report is organized to be read top-to-"
        "bottom: cover → this purpose page → executive summary → AI asset inventory "
        "(grouped by asset type, with one card per registered asset) → recent lifecycle "
        "changes → regulatory coverage table mapping the report sections back to the "
        "canonical regulatory requirements they evidence.",
        font=BODY_FONT, size=11, color=BODY_GRAY, space_after=10,
    )

    _add_styled_para(
        doc,
        "Data lineage",
        font=HEADING_FONT, size=13, bold=True, color=BRAND_BLUE,
        space_before=10, space_after=4,
    )
    _add_styled_para(
        doc,
        "All figures in this report are queried at generation time from the Verity "
        "compliance metamodel — frameworks, provisions, canonical requirements, "
        "Verity features, and the bidirectional bridges that connect them. No data "
        "is imported from external systems; nothing is computed offline. Every value "
        "is traceable back to a source row in the metamodel.",
        font=BODY_FONT, size=11, color=BODY_GRAY,
    )

    _add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # PAGE 3 — EXECUTIVE SUMMARY
    # ════════════════════════════════════════════════════════════
    _add_heading(doc, "Executive Summary", level=1)

    _add_styled_para(
        doc,
        "At the time of generation, Verity governs {{ total_count }} AI asset version(s) "
        "across {{ by_entity_type|length }} asset type(s). The asset inventory is "
        "broken down as follows:",
        space_after=8,
    )
    _add_styled_para(
        doc,
        "• Agents: {{ agent_count }}     • Tasks: {{ task_count }}     "
        "• Prompts: {{ prompt_count }}",
        font=BODY_FONT, size=11, bold=True, color=BRAND_BLUE, space_after=8,
    )
    _add_styled_para(
        doc,
        "By materiality tier: {{ high_count }} high, {{ medium_count }} medium, "
        "{{ low_count }} low. {{ lifecycle_events_total }} lifecycle state "
        "transitions are recorded in the recent-changes section of this report.",
    )

    _add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # PAGE 4 — AI ASSET INVENTORY
    # ════════════════════════════════════════════════════════════
    _add_heading(doc, "AI Asset Inventory", level=1)

    _add_styled_para(
        doc,
        "Every AI asset registered to Verity is listed below, grouped by asset type. "
        "Each card shows the asset's display name, current version label, lifecycle "
        "state, plain-language description, materiality tier, named owner, and the "
        "applications it serves.",
        space_after=14,
    )

    # ── Agents subsection ──────────────────────────────────────
    _add_heading(doc, "Agents ({{ agent_count }})", level=2)
    _add_asset_card_table(doc, item_var="a", list_var="agents")

    # ── Tasks subsection ───────────────────────────────────────
    _add_heading(doc, "Tasks ({{ task_count }})", level=2, page_break_before=False)
    _add_styled_para(doc, "", space_before=8)
    _add_asset_card_table(doc, item_var="a", list_var="tasks")

    # ── Prompts subsection ─────────────────────────────────────
    _add_heading(doc, "Prompts ({{ prompt_count }})", level=2, page_break_before=False)
    _add_styled_para(doc, "", space_before=8)
    _add_asset_card_table(doc, item_var="a", list_var="prompts")

    _add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # PAGE 5 — RECENT LIFECYCLE CHANGES
    # ════════════════════════════════════════════════════════════
    _add_heading(doc, "Recent Lifecycle Changes", level=1)

    _add_styled_para(
        doc,
        "Up to the 50 most recent state transitions across all assets, sub-grouped "
        "by asset type. Each row reflects a HITL approval gate.",
        space_after=10,
    )

    _add_heading(doc, "Agents", level=2)
    _add_lifecycle_table(doc, "lifecycle_agents")

    _add_styled_para(doc, "", space_before=10)
    _add_heading(doc, "Tasks", level=2)
    _add_lifecycle_table(doc, "lifecycle_tasks")

    _add_styled_para(doc, "", space_before=10)
    _add_heading(doc, "Prompts", level=2)
    _add_lifecycle_table(doc, "lifecycle_prompts")

    _add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # PAGE 6 — REGULATORY COVERAGE
    # ════════════════════════════════════════════════════════════
    _add_heading(doc, "Regulatory Coverage", level=1)

    _add_styled_para(
        doc,
        "This report provides evidence for the following canonical regulatory "
        "requirements within the Verity compliance metamodel.",
        space_after=10,
    )

    can_table = doc.add_table(rows=4, cols=4)
    can_table.style = "Light Grid Accent 1"
    headers = ["Theme", "Canonical Requirement", "Coverage", "Section"]
    for i, h in enumerate(headers):
        cell = can_table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        _set_run_font(run, name=BODY_FONT, size=10, bold=True, color=BRAND_BLUE)

    can_table.rows[1].cells[0].text = "{%tr for c in canonicals %}"
    can_table.rows[1].cells[0].merge(can_table.rows[1].cells[3])

    body = can_table.rows[2].cells
    cells_text = [
        "{{ c.theme_name }}",
        "{{ c.title }}",
        "{{ c.coverage_level or '—' }}",
        "{{ c.section or '—' }}",
    ]
    for i, txt in enumerate(cells_text):
        body[i].text = ""
        run = body[i].paragraphs[0].add_run(txt)
        _set_run_font(run, name=BODY_FONT, size=10, color=BODY_GRAY)

    can_table.rows[3].cells[0].text = "{%tr endfor %}"
    can_table.rows[3].cells[0].merge(can_table.rows[3].cells[3])

    # ── Save ───────────────────────────────────────────────────
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    out = TEMPLATES_DIR / "model_inventory.docx"
    doc.save(out)
    return out


# =============================================================================
# Dispatcher
# =============================================================================

# =============================================================================
# Decision Audit Trail
# =============================================================================

def author_workflow_audit_trail() -> Path:
    """End-to-end audit trail for one workflow_run_id."""
    doc = Document()
    _set_normal_style(doc)
    _setup_section(doc)

    # ── COVER ──
    for _ in range(8):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("VERITY"); _set_run_font(run, name=HEADING_FONT, size=14, bold=True, color=BRAND_BLUE_LIGHT)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Compliance Report"); _set_run_font(run, name=HEADING_FONT, size=11, color=MUTED_GRAY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28); p.paragraph_format.space_after = Pt(8)
    run = p.add_run("{{ report_title }}")
    _set_run_font(run, name=HEADING_FONT, size=32, bold=True, color=BRAND_BLUE)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run("Generated {{ generated_at_str }}")
    _set_run_font(run, name=BODY_FONT, size=12, color=BODY_GRAY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Workflow: {{ workflow_run_id }}")
    _set_run_font(run, name=BODY_FONT, size=10, italic=True, color=MUTED_GRAY)
    _add_page_break(doc)

    # ── PURPOSE ──
    _add_heading(doc, "Purpose", level=1)
    _add_styled_para(doc, "{{ report_description }}", space_after=12)
    _add_styled_para(
        doc,
        "This report covers one workflow execution, identified by the "
        "execution_context_id above. Every decision the workflow produced is "
        "listed in chronological order, with full reasoning, confidence "
        "scores, HITL gate status, and any human overrides applied.",
        space_after=10,
    )
    _add_page_break(doc)

    # ── EXECUTIVE SUMMARY ──
    _add_heading(doc, "Executive Summary", level=1)
    _add_styled_para(
        doc,
        "{{ decision_count }} decision(s) were produced by {{ distinct_assets }} "
        "distinct asset(s) over a total runtime of {{ total_duration_ms }} ms. "
        "The mean reported confidence across decisions was "
        "{{ '%.2f'|format(avg_confidence) if avg_confidence is not none else 'n/a' }}. "
        "{{ hitl_required_count }} decision(s) required HITL approval; "
        "{{ override_count }} override(s) were recorded.",
    )
    _add_page_break(doc)

    # ── DECISION SEQUENCE ──
    _add_heading(doc, "Decision Sequence", level=1)
    _add_styled_para(
        doc,
        "Each card below is one decision, ordered chronologically. The card "
        "shows the asset that produced it, the workflow step, the LLM used, "
        "duration, confidence, HITL state, and the reasoning text the model "
        "produced.",
        space_after=10,
    )

    # Decision cards table — single-column shaded cell per decision.
    table = doc.add_table(rows=3, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)

    table.rows[0].cells[0].text = ""
    table.rows[0].cells[0].paragraphs[0].add_run("{%tr for d in decisions %}")

    body = table.rows[1].cells[0]
    body.width = Inches(6.5)
    _shade_cell(body, CARD_SHADE_HEX)
    _set_cell_borders(body, color="C5CDDC")
    body._tc.remove(body.paragraphs[0]._p)

    p = body.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Step: {{ d.step_name or '—' }}")
    _set_run_font(run, name=BODY_FONT, size=12, bold=True, color=BRAND_BLUE)
    run = p.add_run("    ·    ")
    _set_run_font(run, name=BODY_FONT, size=11, color=MUTED_GRAY)
    run = p.add_run("{{ d.asset_type_display }}: {{ d.asset_display_name }} v{{ d.version_label }}")
    _set_run_font(run, name=BODY_FONT, size=11, color=BODY_GRAY)

    p = body.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
    run = p.add_run("{{ d.created_at.strftime('%Y-%m-%d %H:%M:%S') if d.created_at else '—' }}")
    _set_run_font(run, name=BODY_FONT, size=9, italic=True, color=MUTED_GRAY)
    run = p.add_run("    ·    Model: {{ d.model_used or '—' }}    ·    Duration: {{ d.duration_ms or 0 }} ms    ·    Confidence: {{ '%.2f'|format(d.confidence_score|float) if d.confidence_score is not none else '—' }}")
    _set_run_font(run, name=BODY_FONT, size=9, color=MUTED_GRAY)

    p = body.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run("Input: ")
    _set_run_font(run, name=BODY_FONT, size=10, bold=True, color=BODY_GRAY)
    run = p.add_run("{{ d.input_summary or '—' }}")
    _set_run_font(run, name=BODY_FONT, size=10, color=BODY_GRAY)

    p = body.add_paragraph()
    run = p.add_run("Output: ")
    _set_run_font(run, name=BODY_FONT, size=10, bold=True, color=BODY_GRAY)
    run = p.add_run("{{ d.output_summary or '—' }}")
    _set_run_font(run, name=BODY_FONT, size=10, color=BODY_GRAY)

    p = body.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("Reasoning: ")
    _set_run_font(run, name=BODY_FONT, size=10, bold=True, color=BODY_GRAY)
    run = p.add_run("{{ d.reasoning_text or '—' }}")
    _set_run_font(run, name=BODY_FONT, size=10, italic=True, color=BODY_GRAY)

    table.rows[2].cells[0].text = "{%tr endfor %}"

    _add_page_break(doc)

    # ── HITL OVERRIDES ──
    _add_heading(doc, "HITL Overrides", level=1)
    _add_styled_para(
        doc,
        "{{ override_count }} override(s) were applied to decisions in this "
        "workflow. Each override represents a named reviewer correcting an "
        "AI-produced value before it became binding.",
        space_after=10,
    )
    ov_table = doc.add_table(rows=4, cols=5)
    ov_table.style = "Light Grid Accent 1"
    headers = ["Fact", "AI Value", "Human Value", "Reviewer", "When"]
    for i, h in enumerate(headers):
        cell = ov_table.rows[0].cells[i]; cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        _set_run_font(run, name=BODY_FONT, size=10, bold=True, color=BRAND_BLUE)
    ov_table.rows[1].cells[0].text = "{%tr for o in overrides %}"
    ov_table.rows[1].cells[0].merge(ov_table.rows[1].cells[4])
    body = ov_table.rows[2].cells
    cells_text = [
        "{{ o.fact_type }}",
        "{{ o.ai_value or '—' }}",
        "{{ o.hitl_value }}",
        "{{ o.overridden_by }}",
        "{{ o.created_at.strftime('%Y-%m-%d %H:%M') if o.created_at else '—' }}",
    ]
    for i, txt in enumerate(cells_text):
        body[i].text = ""
        run = body[i].paragraphs[0].add_run(txt)
        _set_run_font(run, name=BODY_FONT, size=10, color=BODY_GRAY)
    ov_table.rows[3].cells[0].text = "{%tr endfor %}"
    ov_table.rows[3].cells[0].merge(ov_table.rows[3].cells[4])

    _add_page_break(doc)

    # ── REGULATORY COVERAGE ──
    _add_heading(doc, "Regulatory Coverage", level=1)
    _add_canonicals_table(doc)

    out = TEMPLATES_DIR / "workflow_audit_trail.docx"
    doc.save(out)
    return out


# =============================================================================
# Decision Audit Trail — single-decision deep-dive
# =============================================================================

def author_decision_audit_trail() -> Path:
    """Single-decision report. Required scope: decision_id."""
    doc = Document()
    _set_normal_style(doc)
    _setup_section(doc)

    # ── COVER ──
    for _ in range(8):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("VERITY"); _set_run_font(run, name=HEADING_FONT, size=14, bold=True, color=BRAND_BLUE_LIGHT)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run("Compliance Report"); _set_run_font(run, name=HEADING_FONT, size=11, color=MUTED_GRAY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    run = p.add_run("{{ report_title }}")
    _set_run_font(run, name=HEADING_FONT, size=30, bold=True, color=BRAND_BLUE)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run("{{ decision.asset_type_display }}: {{ decision.asset_display_name }} v{{ decision.version_label }}")
    _set_run_font(run, name=HEADING_FONT, size=15, bold=True, color=BRAND_BLUE_LIGHT)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Decision rendered {{ decision.created_at.strftime('%Y-%m-%d %H:%M:%S') if decision.created_at else '—' }}")
    _set_run_font(run, name=BODY_FONT, size=11, color=BODY_GRAY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run("Generated {{ generated_at_str }}")
    _set_run_font(run, name=BODY_FONT, size=10, italic=True, color=MUTED_GRAY)
    _add_page_break(doc)

    # ── PURPOSE ──
    _add_heading(doc, "Purpose", level=1)
    _add_styled_para(doc, "{{ report_description }}", space_after=12)
    _add_page_break(doc)

    # ── DECISION CONTEXT ──
    _add_heading(doc, "Decision Context", level=1)
    grid = doc.add_table(rows=10, cols=2)
    grid.style = "Light Grid Accent 1"
    rows = [
        ("Asset",           "{{ decision.asset_type_display }}: {{ decision.asset_display_name }}"),
        ("Version",         "v{{ decision.version_label }}"),
        ("Owner",           "{{ decision.owner_name }}"),
        ("Materiality",     "{{ decision.materiality_display }}"),
        ("Workflow step",   "{{ decision.step_name or '—' }}"),
        ("Application",     "{{ decision.application_code or '—' }}"),
        ("Channel",         "{{ decision.channel or '—' }}"),
        ("Run purpose",     "{{ decision.run_purpose or '—' }}"),
        ("LLM used",        "{{ decision.model_used or '—' }}"),
        ("Decision ID",     "{{ decision.decision_id }}"),
    ]
    for i, (k, v) in enumerate(rows):
        c0 = grid.rows[i].cells[0]; c1 = grid.rows[i].cells[1]
        c0.text = ""; run = c0.paragraphs[0].add_run(k)
        _set_run_font(run, name=BODY_FONT, size=10, bold=True, color=BRAND_BLUE)
        c1.text = ""; run = c1.paragraphs[0].add_run(v)
        _set_run_font(run, name=BODY_FONT, size=10, color=BODY_GRAY)
    _add_styled_para(doc, "", space_before=8)
    _add_styled_para(doc, "Asset description", font=HEADING_FONT, size=12, bold=True,
                     color=BRAND_BLUE, space_before=10, space_after=2)
    _add_styled_para(doc, "{{ decision.asset_description or '—' }}",
                     italic=True, color=MUTED_GRAY)
    _add_page_break(doc)

    # ── REASONING ──
    _add_heading(doc, "Reasoning", level=1)
    _add_styled_para(
        doc,
        "Below is the model-emitted reasoning for this decision, exactly as "
        "produced at runtime. This is the substantive justification a "
        "regulator or affected consumer is entitled to see.",
        space_after=10,
    )
    _add_styled_para(doc, "Input summary", font=HEADING_FONT, size=12, bold=True,
                     color=BRAND_BLUE, space_before=6, space_after=2)
    _add_styled_para(doc, "{{ decision.input_summary or '—' }}", color=BODY_GRAY)

    _add_styled_para(doc, "Output summary", font=HEADING_FONT, size=12, bold=True,
                     color=BRAND_BLUE, space_before=14, space_after=2)
    _add_styled_para(doc, "{{ decision.output_summary or '—' }}", color=BODY_GRAY)

    _add_styled_para(doc, "Reasoning", font=HEADING_FONT, size=12, bold=True,
                     color=BRAND_BLUE, space_before=14, space_after=2)
    _add_styled_para(doc, "{{ decision.reasoning_text or '—' }}",
                     italic=True, color=BODY_GRAY)
    _add_page_break(doc)

    # ── METRICS ──
    _add_heading(doc, "Confidence & Performance", level=1)
    grid = doc.add_table(rows=4, cols=2)
    grid.style = "Light Grid Accent 1"
    rows = [
        ("Confidence",        "{{ '%.2f'|format(decision.confidence_score|float) if decision.confidence_score is not none else '—' }}"),
        ("Duration",          "{{ decision.duration_ms or 0 }} ms"),
        ("Tokens (in/out)",   "{{ decision.input_tokens or 0 }} / {{ decision.output_tokens or 0 }}"),
        ("HITL gate",         "{% if decision.hitl_required %}required{% if decision.hitl_completed %} (completed){% endif %}{% else %}not required{% endif %}"),
    ]
    for i, (k, v) in enumerate(rows):
        c0 = grid.rows[i].cells[0]; c1 = grid.rows[i].cells[1]
        c0.text = ""; run = c0.paragraphs[0].add_run(k)
        _set_run_font(run, name=BODY_FONT, size=10, bold=True, color=BRAND_BLUE)
        c1.text = ""; run = c1.paragraphs[0].add_run(v)
        _set_run_font(run, name=BODY_FONT, size=10, color=BODY_GRAY)
    _add_page_break(doc)

    # ── OVERRIDES ──
    _add_heading(doc, "Human Overrides", level=1)
    _add_styled_para(
        doc,
        "{{ override_count }} human override(s) applied to this specific decision. "
        "Each row records a named reviewer correcting an AI-produced value.",
        space_after=10,
    )
    ov = doc.add_table(rows=4, cols=5)
    ov.style = "Light Grid Accent 1"
    headers = ["Fact", "AI Value", "Human Value", "Reviewer", "When"]
    for i, h in enumerate(headers):
        cell = ov.rows[0].cells[i]; cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        _set_run_font(run, name=BODY_FONT, size=10, bold=True, color=BRAND_BLUE)
    ov.rows[1].cells[0].text = "{%tr for o in overrides %}"
    ov.rows[1].cells[0].merge(ov.rows[1].cells[4])
    body = ov.rows[2].cells
    cells_text = [
        "{{ o.fact_type }}",
        "{{ o.ai_value or '—' }}",
        "{{ o.hitl_value }}",
        "{{ o.overridden_by }}",
        "{{ o.created_at.strftime('%Y-%m-%d %H:%M') if o.created_at else '—' }}",
    ]
    for i, txt in enumerate(cells_text):
        body[i].text = ""
        run = body[i].paragraphs[0].add_run(txt)
        _set_run_font(run, name=BODY_FONT, size=10, color=BODY_GRAY)
    ov.rows[3].cells[0].text = "{%tr endfor %}"
    ov.rows[3].cells[0].merge(ov.rows[3].cells[4])
    _add_page_break(doc)

    # ── REGULATORY COVERAGE ──
    _add_heading(doc, "Regulatory Coverage", level=1)
    _add_canonicals_table(doc)

    out = TEMPLATES_DIR / "decision_audit_trail.docx"
    doc.save(out)
    return out


# =============================================================================
# Fairness Validation Summary
# =============================================================================

def author_fairness_validation_summary() -> Path:
    doc = Document()
    _set_normal_style(doc)
    _setup_section(doc)

    # ── COVER ──
    for _ in range(8):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("VERITY"); _set_run_font(run, name=HEADING_FONT, size=14, bold=True, color=BRAND_BLUE_LIGHT)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run("Compliance Report"); _set_run_font(run, name=HEADING_FONT, size=11, color=MUTED_GRAY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    run = p.add_run("{{ report_title }}")
    _set_run_font(run, name=HEADING_FONT, size=32, bold=True, color=BRAND_BLUE)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run("Generated {{ generated_at_str }}")
    _set_run_font(run, name=BODY_FONT, size=12, color=BODY_GRAY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Scope: {{ scope_summary }}")
    _set_run_font(run, name=BODY_FONT, size=10, italic=True, color=MUTED_GRAY)
    _add_page_break(doc)

    # ── PURPOSE ──
    _add_heading(doc, "Purpose", level=1)
    _add_styled_para(doc, "{{ report_description }}", space_after=12)
    _add_page_break(doc)

    # ── EXECUTIVE SUMMARY ──
    _add_heading(doc, "Executive Summary", level=1)
    _add_styled_para(
        doc,
        "{{ result_count }} validation result(s) recorded across {{ asset_count }} "
        "asset(s). {{ pass_count }} passed, {{ fail_count }} failed.",
    )
    _add_page_break(doc)

    # ── PER-ASSET VALIDATION ──
    _add_heading(doc, "Validation by Asset", level=1)
    _add_styled_para(
        doc,
        "Each asset is summarized below with its registered metadata and the "
        "list of validation results recorded against it.",
        space_after=10,
    )

    table = doc.add_table(rows=3, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)

    table.rows[0].cells[0].text = "{%tr for asset in asset_summaries %}"

    body = table.rows[1].cells[0]
    _shade_cell(body, CARD_SHADE_HEX)
    _set_cell_borders(body, color="C5CDDC")
    body._tc.remove(body.paragraphs[0]._p)

    p = body.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    run = p.add_run("{{ asset.asset_type_display }}: {{ asset.asset_display_name }} v{{ asset.version_label }}")
    _set_run_font(run, name=BODY_FONT, size=12, bold=True, color=BRAND_BLUE)

    p = body.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Materiality: {{ asset.materiality_display }}    ·    Owner: {{ asset.owner_name }}    ·    {{ asset.pass_count }} passed, {{ asset.fail_count }} failed")
    _set_run_font(run, name=BODY_FONT, size=10, color=MUTED_GRAY)

    # Inner: per-result details (run inside the same shaded cell)
    p = body.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("{% for r in asset.results %}{{ r.metric_type }} — {{ r.passed_display }} (run {{ r.run_at.strftime('%Y-%m-%d') if r.run_at else '—' }}){% if not loop.last %}; {% endif %}{% endfor %}")
    _set_run_font(run, name=BODY_FONT, size=9, italic=True, color=BODY_GRAY)

    table.rows[2].cells[0].text = "{%tr endfor %}"

    _add_page_break(doc)

    # ── DETAILED RESULTS ──
    _add_heading(doc, "Detailed Results", level=1)
    _add_styled_para(
        doc,
        "Every individual validation row, in reverse-chronological order.",
        space_after=10,
    )
    rows_table = doc.add_table(rows=4, cols=6)
    rows_table.style = "Light Grid Accent 1"
    headers = ["Asset", "Suite", "Metric", "Result", "When", "Channel"]
    for i, h in enumerate(headers):
        cell = rows_table.rows[0].cells[i]; cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        _set_run_font(run, name=BODY_FONT, size=10, bold=True, color=BRAND_BLUE)
    rows_table.rows[1].cells[0].text = "{%tr for r in validation_results %}"
    rows_table.rows[1].cells[0].merge(rows_table.rows[1].cells[5])
    body = rows_table.rows[2].cells
    cells_text = [
        "{{ r.asset_display_name }}",
        "{{ r.suite_id }}",
        "{{ r.metric_type or '—' }}",
        "{{ r.passed_display }}",
        "{{ r.run_at.strftime('%Y-%m-%d %H:%M') if r.run_at else '—' }}",
        "{{ r.channel or '—' }}",
    ]
    for i, txt in enumerate(cells_text):
        body[i].text = ""
        run = body[i].paragraphs[0].add_run(txt)
        _set_run_font(run, name=BODY_FONT, size=9, color=BODY_GRAY)
    rows_table.rows[3].cells[0].text = "{%tr endfor %}"
    rows_table.rows[3].cells[0].merge(rows_table.rows[3].cells[5])

    _add_page_break(doc)
    _add_heading(doc, "Regulatory Coverage", level=1)
    _add_canonicals_table(doc)

    out = TEMPLATES_DIR / "fairness_validation_summary.docx"
    doc.save(out)
    return out


# =============================================================================
# NAIC Exhibit C — High-Risk System Deep Dive
# =============================================================================

def author_naic_exhibit_c() -> Path:
    doc = Document()
    _set_normal_style(doc)
    _setup_section(doc)

    # ── COVER ──
    for _ in range(8):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("VERITY"); _set_run_font(run, name=HEADING_FONT, size=14, bold=True, color=BRAND_BLUE_LIGHT)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run("Regulatory Examination Exhibit"); _set_run_font(run, name=HEADING_FONT, size=11, color=MUTED_GRAY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    run = p.add_run("{{ report_title }}")
    _set_run_font(run, name=HEADING_FONT, size=28, bold=True, color=BRAND_BLUE)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    run = p.add_run("{{ target.asset_type_display }}: {{ target.display_name }}")
    _set_run_font(run, name=HEADING_FONT, size=18, bold=True, color=BRAND_BLUE_LIGHT)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Version {{ target.version_label }}    ·    {{ target.lifecycle_state_display }}")
    _set_run_font(run, name=BODY_FONT, size=12, color=BODY_GRAY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run("Generated {{ generated_at_str }}")
    _set_run_font(run, name=BODY_FONT, size=11, color=BODY_GRAY)
    _add_page_break(doc)

    # ── PURPOSE ──
    _add_heading(doc, "Purpose and Audience", level=1)
    _add_styled_para(doc, "{{ report_description }}", space_after=12)
    _add_styled_para(
        doc,
        "This exhibit assembles the complete governance trail for one "
        "high-risk AI system — registration, ownership, lifecycle, validation, "
        "every decision the asset has made in production, and every override "
        "applied to those decisions. Designed to be the single document handed "
        "to a market-conduct examiner.",
    )
    _add_page_break(doc)

    # ── ASSET IDENTIFICATION ──
    _add_heading(doc, "Asset Identification", level=1)
    grid = doc.add_table(rows=8, cols=2)
    grid.style = "Light Grid Accent 1"
    rows = [
        ("Type",          "{{ target.asset_type_display }}"),
        ("Display name",  "{{ target.display_name }}"),
        ("Code name",     "{{ target.entity_name }}"),
        ("Version",       "v{{ target.version_label }}"),
        ("Lifecycle",     "{{ target.lifecycle_state_display }}"),
        ("Materiality",   "{{ target.materiality_display }}"),
        ("Domain",        "{{ target.domain or '—' }}"),
        ("Application(s)", "{{ target.applications }}"),
    ]
    for i, (k, v) in enumerate(rows):
        c0 = grid.rows[i].cells[0]; c1 = grid.rows[i].cells[1]
        c0.text = ""; run = c0.paragraphs[0].add_run(k)
        _set_run_font(run, name=BODY_FONT, size=10, bold=True, color=BRAND_BLUE)
        c1.text = ""; run = c1.paragraphs[0].add_run(v)
        _set_run_font(run, name=BODY_FONT, size=10, color=BODY_GRAY)
    _add_styled_para(doc, "", space_before=8)
    _add_styled_para(doc, "Description", font=HEADING_FONT, size=12, bold=True, color=BRAND_BLUE, space_before=10, space_after=2)
    _add_styled_para(doc, "{{ target.entity_description or '—' }}", italic=True, color=MUTED_GRAY)
    _add_page_break(doc)

    # ── OWNERSHIP & ACCOUNTABILITY ──
    _add_heading(doc, "Ownership & Accountability", level=1)
    grid = doc.add_table(rows=2, cols=2)
    grid.style = "Light Grid Accent 1"
    rows = [
        ("Owner name",  "{{ target.owner_name or '—' }}"),
        ("Owner email", "{{ target.owner_email or '—' }}"),
    ]
    for i, (k, v) in enumerate(rows):
        c0 = grid.rows[i].cells[0]; c1 = grid.rows[i].cells[1]
        c0.text = ""; run = c0.paragraphs[0].add_run(k)
        _set_run_font(run, name=BODY_FONT, size=10, bold=True, color=BRAND_BLUE)
        c1.text = ""; run = c1.paragraphs[0].add_run(v)
        _set_run_font(run, name=BODY_FONT, size=10, color=BODY_GRAY)
    _add_page_break(doc)

    # ── LIFECYCLE HISTORY ──
    _add_heading(doc, "Lifecycle History", level=1)
    _add_styled_para(
        doc,
        "{{ lifecycle_count }} HITL approval gate(s) recorded for this asset.",
        space_after=8,
    )
    lc_table = doc.add_table(rows=4, cols=5)
    lc_table.style = "Light Grid Accent 1"
    headers = ["Gate", "Transition", "Approver", "Role", "When"]
    for i, h in enumerate(headers):
        cell = lc_table.rows[0].cells[i]; cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        _set_run_font(run, name=BODY_FONT, size=10, bold=True, color=BRAND_BLUE)
    lc_table.rows[1].cells[0].text = "{%tr for ev in lifecycle_events %}"
    lc_table.rows[1].cells[0].merge(lc_table.rows[1].cells[4])
    body = lc_table.rows[2].cells
    cells_text = [
        "{{ ev.gate_type or '—' }}",
        "{{ ev.from_state_display }} → {{ ev.to_state_display }}",
        "{{ ev.approver_name or '—' }}",
        "{{ ev.approver_role or '—' }}",
        "{{ ev.approved_at.strftime('%Y-%m-%d %H:%M') if ev.approved_at else '—' }}",
    ]
    for i, txt in enumerate(cells_text):
        body[i].text = ""
        run = body[i].paragraphs[0].add_run(txt)
        _set_run_font(run, name=BODY_FONT, size=10, color=BODY_GRAY)
    lc_table.rows[3].cells[0].text = "{%tr endfor %}"
    lc_table.rows[3].cells[0].merge(lc_table.rows[3].cells[4])
    _add_page_break(doc)

    # ── VALIDATION & TESTING ──
    _add_heading(doc, "Validation & Testing", level=1)
    _add_styled_para(
        doc,
        "{{ validation_count }} validation result(s) recorded. "
        "{{ validation_pass_count }} passed, {{ validation_fail_count }} failed.",
        space_after=8,
    )
    vt = doc.add_table(rows=4, cols=5)
    vt.style = "Light Grid Accent 1"
    headers = ["Suite", "Metric", "Result", "When", "Channel"]
    for i, h in enumerate(headers):
        cell = vt.rows[0].cells[i]; cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        _set_run_font(run, name=BODY_FONT, size=10, bold=True, color=BRAND_BLUE)
    vt.rows[1].cells[0].text = "{%tr for v in validation_results %}"
    vt.rows[1].cells[0].merge(vt.rows[1].cells[4])
    body = vt.rows[2].cells
    cells_text = [
        "{{ v.suite_id }}",
        "{{ v.metric_type or '—' }}",
        "{{ v.passed_display }}",
        "{{ v.run_at.strftime('%Y-%m-%d %H:%M') if v.run_at else '—' }}",
        "{{ v.channel or '—' }}",
    ]
    for i, txt in enumerate(cells_text):
        body[i].text = ""
        run = body[i].paragraphs[0].add_run(txt)
        _set_run_font(run, name=BODY_FONT, size=10, color=BODY_GRAY)
    vt.rows[3].cells[0].text = "{%tr endfor %}"
    vt.rows[3].cells[0].merge(vt.rows[3].cells[4])
    _add_page_break(doc)

    # ── DECISIONS & PRODUCTION MONITORING ──
    _add_heading(doc, "Production Decisions", level=1)
    _add_styled_para(
        doc,
        "{{ decision_count }} most recent decision(s) produced by this asset. "
        "Mean confidence: {{ '%.2f'|format(avg_confidence) if avg_confidence is not none else 'n/a' }}.",
        space_after=8,
    )
    dt = doc.add_table(rows=4, cols=6)
    dt.style = "Light Grid Accent 1"
    headers = ["When", "Step", "Model", "Confidence", "Duration (ms)", "HITL"]
    for i, h in enumerate(headers):
        cell = dt.rows[0].cells[i]; cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        _set_run_font(run, name=BODY_FONT, size=10, bold=True, color=BRAND_BLUE)
    dt.rows[1].cells[0].text = "{%tr for d in decisions %}"
    dt.rows[1].cells[0].merge(dt.rows[1].cells[5])
    body = dt.rows[2].cells
    cells_text = [
        "{{ d.created_at.strftime('%Y-%m-%d %H:%M') if d.created_at else '—' }}",
        "{{ d.step_name or '—' }}",
        "{{ d.model_used or '—' }}",
        "{{ '%.2f'|format(d.confidence_score|float) if d.confidence_score is not none else '—' }}",
        "{{ d.duration_ms or 0 }}",
        "{% if d.hitl_required %}required{% if d.hitl_completed %} (done){% endif %}{% else %}—{% endif %}",
    ]
    for i, txt in enumerate(cells_text):
        body[i].text = ""
        run = body[i].paragraphs[0].add_run(txt)
        _set_run_font(run, name=BODY_FONT, size=9, color=BODY_GRAY)
    dt.rows[3].cells[0].text = "{%tr endfor %}"
    dt.rows[3].cells[0].merge(dt.rows[3].cells[5])
    _add_page_break(doc)

    # ── HITL OVERRIDES ──
    _add_heading(doc, "HITL & Overrides", level=1)
    _add_styled_para(
        doc,
        "{{ override_count }} human override(s) recorded against this asset's "
        "decisions.",
        space_after=8,
    )
    ot = doc.add_table(rows=4, cols=5)
    ot.style = "Light Grid Accent 1"
    headers = ["Fact", "AI Value", "Human Value", "Reviewer", "When"]
    for i, h in enumerate(headers):
        cell = ot.rows[0].cells[i]; cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        _set_run_font(run, name=BODY_FONT, size=10, bold=True, color=BRAND_BLUE)
    ot.rows[1].cells[0].text = "{%tr for o in overrides %}"
    ot.rows[1].cells[0].merge(ot.rows[1].cells[4])
    body = ot.rows[2].cells
    cells_text = [
        "{{ o.fact_type }}",
        "{{ o.ai_value or '—' }}",
        "{{ o.hitl_value }}",
        "{{ o.overridden_by }}",
        "{{ o.created_at.strftime('%Y-%m-%d %H:%M') if o.created_at else '—' }}",
    ]
    for i, txt in enumerate(cells_text):
        body[i].text = ""
        run = body[i].paragraphs[0].add_run(txt)
        _set_run_font(run, name=BODY_FONT, size=10, color=BODY_GRAY)
    ot.rows[3].cells[0].text = "{%tr endfor %}"
    ot.rows[3].cells[0].merge(ot.rows[3].cells[4])
    _add_page_break(doc)

    # ── REGULATORY COVERAGE ──
    _add_heading(doc, "Regulatory Coverage", level=1)
    _add_canonicals_table(doc)

    out = TEMPLATES_DIR / "naic_exhibit_c.docx"
    doc.save(out)
    return out


# =============================================================================
# Shared helpers extracted for the new templates
# =============================================================================

def _set_normal_style(doc):
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    rPr = normal.element.get_or_add_rPr()
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), BODY_FONT)
    rFonts.set(qn("w:hAnsi"), BODY_FONT)
    rFonts.set(qn("w:cs"),    BODY_FONT)
    rFonts.set(qn("w:eastAsia"), BODY_FONT)
    rPr.append(rFonts)


def _setup_section(doc):
    section = doc.sections[0]
    section.top_margin    = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin   = Inches(0.9)
    section.right_margin  = Inches(0.9)
    _enable_different_first_page_header_footer(section)
    fp_header = section.first_page_header; fp_header.is_linked_to_previous = False
    fp_header.paragraphs[0].text = ""
    fp_footer = section.first_page_footer; fp_footer.is_linked_to_previous = False
    fp_footer.paragraphs[0].text = ""
    header_p = section.header.paragraphs[0]; header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_p.add_run("VERITY · {{ report_title }}")
    _set_run_font(run, name=HEADING_FONT, size=9, bold=True, color=BRAND_BLUE)
    footer_p = section.footer.paragraphs[0]; footer_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = footer_p.add_run("{{ report_title }}  ·  Generated {{ generated_at_str }}")
    _set_run_font(run, name=BODY_FONT, size=8, color=MUTED_GRAY)


def _add_canonicals_table(doc):
    can_table = doc.add_table(rows=4, cols=4)
    can_table.style = "Light Grid Accent 1"
    headers = ["Theme", "Canonical Requirement", "Coverage", "Section"]
    for i, h in enumerate(headers):
        cell = can_table.rows[0].cells[i]; cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        _set_run_font(run, name=BODY_FONT, size=10, bold=True, color=BRAND_BLUE)
    can_table.rows[1].cells[0].text = "{%tr for c in canonicals %}"
    can_table.rows[1].cells[0].merge(can_table.rows[1].cells[3])
    body = can_table.rows[2].cells
    cells_text = [
        "{{ c.theme_name }}",
        "{{ c.title }}",
        "{{ c.coverage_level or '—' }}",
        "{{ c.section or '—' }}",
    ]
    for i, txt in enumerate(cells_text):
        body[i].text = ""
        run = body[i].paragraphs[0].add_run(txt)
        _set_run_font(run, name=BODY_FONT, size=10, color=BODY_GRAY)
    can_table.rows[3].cells[0].text = "{%tr endfor %}"
    can_table.rows[3].cells[0].merge(can_table.rows[3].cells[3])


AUTHORS = {
    "model_inventory":             author_model_inventory,
    "decision_audit_trail":        author_decision_audit_trail,    # single decision
    "workflow_audit_trail":        author_workflow_audit_trail,    # one workflow
    "fairness_validation_summary": author_fairness_validation_summary,
    "naic_exhibit_c":              author_naic_exhibit_c,
}


def main() -> None:
    if len(sys.argv) < 2:
        print(f"Usage: python -m verity.reporting._template_authoring <report_code>")
        print(f"Available: {list(AUTHORS)}")
        sys.exit(1)
    code = sys.argv[1]
    if code not in AUTHORS:
        print(f"Unknown report {code!r}. Available: {list(AUTHORS)}")
        sys.exit(1)
    out = AUTHORS[code]()
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
